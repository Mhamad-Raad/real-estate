"""Build and register the two placeholder `.docx` letter templates (§6.6, §6.8).

These stand in until the office supplies its real Word files. The **structure** is what matters
and is faithful: the same placeholder names, the same repeating table row, the spouse columns
kept present-but-blank for an unmarried beneficiary, and the list letter's table on its own page.
The cosmetics (letterhead, logo, fonts, CC lists) are deliberately plain.

Swapping in a real template is an upload through the admin screen — no code changes — because the
context contract lives in `documents/letters.py`, not in the file.

    python manage.py build_placeholder_templates
"""

import tempfile
from pathlib import Path

from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management.base import BaseCommand
from docx import Document as Docx
from docx.shared import Pt

from accounts.models import User
from documents.docx_rtl import rtl_paragraph, rtl_table
from documents.models import DocumentTemplate
from documents.services import create_template

# Person columns then spouse columns — one row carries both, so the two tables in the real form
# stay aligned row-for-row when it replaces this placeholder.
COLUMNS = [
    ("ز", "n"),
    ("ناو", "full_name"),
    ("لەدایکبوون", "year"),
    ("ناوی دایک", "mother_name"),
    ("هاوسەر", "spouse_name"),
    ("لەدایکبوون", "spouse_year"),
    ("ناوی دایک", "spouse_mother_name"),
]

LETTERHEAD = [
    "هەرێمی کوردستانی عێراق — ئەنجومەنی وەزیران",
    "وەزارەتی شارەوانی و گەشتوگوزار — سەرۆکایەتیی شارەوانیی سلێمانی",
    "بەشی مولکایەتی و زەوی و زار — هۆبەی زەوی و زار",
]

BODY = (
    "بفەرموون بە ئاگادارکردنەوەمان لە سوودمەندبوونی ئەم بەڕێزانەی خوارەوە و هاوسەرەکانیان لە "
    "وەرگرتنی زەوی یان یەکەی نیشتەجێبوون بە هەر هۆکارێک جگە لە ڕێگەی زیادکردنی ئاشکرا، لەگەڵ "
    "ئاگادارکردنەوەمان لە هۆکار و بنەمای سوودمەندبوونیان:"
)

# The names in this sentence are placeholders, filled from the first and last rows of the table.
LIST_BODY = (
    "بفەرموون بە ئاگادارکردنەوەمان بە وردبینی کردنی ناوی لیستی هاوپێچ کە بە ناوی "
    "({{ first_name }}) دەست پێدەکات و بە ناوی ({{ last_name }}) کۆتایی دێت، کە ئایا "
    "سوودمەندبوون لە وەرگرتنی زەوی یان یەکەی نیشتەجێبوون بە هەر ڕێگەیەک جگە لە ڕێگەی "
    "زیادکردنی ئاشکرا. کۆی ژمارەی ناوەکان: {{ count }}."
)


def _letterhead(doc) -> None:
    for line in LETTERHEAD:
        rtl_paragraph(doc.add_paragraph(), line, size_pt=11)
    # Left blank on purpose: the office writes the number and date by hand.
    rtl_paragraph(doc.add_paragraph(), "ژمارە: ____________", size_pt=11)
    rtl_paragraph(doc.add_paragraph(), "ڕێکەوت: ____________", size_pt=11)


def _beneficiary_table(doc) -> None:
    """Header row, then the loop markers in rows of their own — a {%tr %} deletes its own row."""
    table = rtl_table(doc.add_table(rows=4, cols=len(COLUMNS)))
    table.style = "Table Grid"
    for cell, (heading, _field) in zip(table.rows[0].cells, COLUMNS):
        rtl_paragraph(cell.paragraphs[0], heading, size_pt=10)
    rtl_paragraph(table.rows[1].cells[0].paragraphs[0], "{%tr for r in rows %}")
    for cell, (_heading, field) in zip(table.rows[2].cells, COLUMNS):
        rtl_paragraph(cell.paragraphs[0], f"{{{{ r.{field} }}}}", size_pt=10)
    rtl_paragraph(table.rows[3].cells[0].paragraphs[0], "{%tr endfor %}")


def _closing(doc) -> None:
    rtl_paragraph(doc.add_paragraph(), "لەگەڵ ڕێزدا...", size_pt=11)
    doc.add_paragraph()
    rtl_paragraph(doc.add_paragraph(), "سەرۆکی شارەوانیی سلێمانی", size_pt=11)
    rtl_paragraph(doc.add_paragraph(), "هاوپێچ: وێنەیەک لە بەڵگەنامەکانی ناساندن.", size_pt=10)


def build_eligibility_single(path: Path) -> None:
    doc = Docx()
    doc.styles["Normal"].font.size = Pt(11)
    _letterhead(doc)
    rtl_paragraph(doc.add_paragraph(), "بۆ/ بەڕێوەبەرایەتی گشتی شارەوانییەکانی سلێمانی", size_pt=12)
    rtl_paragraph(doc.add_paragraph(), "بابەت/ سۆراغکردنی سوودمەندی", size_pt=12)
    rtl_paragraph(doc.add_paragraph(), BODY, size_pt=11)
    _beneficiary_table(doc)
    _closing(doc)
    doc.save(path)


def build_process_list(path: Path) -> None:
    doc = Docx()
    doc.styles["Normal"].font.size = Pt(11)
    _letterhead(doc)
    rtl_paragraph(doc.add_paragraph(), "بۆ/ بەڕێوەبەرایەتی تۆمارکردنی خانووبەرەی یەکەمی سلێمانی", size_pt=12)
    rtl_paragraph(doc.add_paragraph(), "بۆ/ بەڕێوەبەرایەتی تۆمارکردنی خانووبەرەی دووەمی سلێمانی", size_pt=12)
    rtl_paragraph(doc.add_paragraph(), "بابەت/ سۆراغکردنی سوودمەندی", size_pt=12)
    rtl_paragraph(doc.add_paragraph(), LIST_BODY, size_pt=11)
    _closing(doc)
    # The list can run to many rows, so the table gets its own page rather than sitting under
    # the letter — the layout the office asked for.
    doc.add_page_break()
    rtl_paragraph(doc.add_paragraph(), "لیستی سوودمەندان", size_pt=12)
    _beneficiary_table(doc)
    doc.save(path)


# Cover-sheet fields: (heading, context key). Rendered as a two-column label/value table.
SUMMARY_FIELDS = [
    ("ناوی سوودمەند", "client_name"),
    ("ژمارەی ناسنامە", "client_pid"),
    ("ناوی دایک", "mother_name"),
    ("ساڵی لەدایکبوون", "birth_year"),
    ("هاوسەر", "spouse_name"),
    ("ژمارەی زەوی", "land_id"),
    ("ناونیشانی زەوی", "land_address"),
    ("پۆل", "category"),
    ("پارێزەر", "assigned_lawyer"),
    ("دۆخی گشتی", "overall_status"),
    ("بەرواری کردنەوە", "created_at"),
    ("ژمارەی بەڵگەنامەکان", "document_count"),
]


def _summary_fields_table(doc) -> None:
    table = rtl_table(doc.add_table(rows=len(SUMMARY_FIELDS), cols=2))
    table.style = "Table Grid"
    for row, (heading, field) in zip(table.rows, SUMMARY_FIELDS):
        rtl_paragraph(row.cells[0].paragraphs[0], heading, size_pt=10)
        rtl_paragraph(row.cells[1].paragraphs[0], f"{{{{ {field} }}}}", size_pt=10)


def _steps_table(doc) -> None:
    headings = ["هەنگاو", "دۆخ", "بەرواری دەستپێک", "بەرواری کۆتایی"]
    fields = ["n", "status", "start_date", "end_date"]
    table = rtl_table(doc.add_table(rows=4, cols=len(headings)))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, headings):
        rtl_paragraph(cell.paragraphs[0], heading, size_pt=10)
    rtl_paragraph(table.rows[1].cells[0].paragraphs[0], "{%tr for s in steps %}")
    for cell, field in zip(table.rows[2].cells, fields):
        rtl_paragraph(cell.paragraphs[0], f"{{{{ s.{field} }}}}", size_pt=10)
    rtl_paragraph(table.rows[3].cells[0].paragraphs[0], "{%tr endfor %}")


def _institutes_table(doc) -> None:
    headings = ["هەنگاو", "دەزگا", "دۆخی پەسەندکردن", "بەروار", "پارێزەر"]
    fields = ["step", "name", "approval_status", "approval_date", "lawyer"]
    table = rtl_table(doc.add_table(rows=4, cols=len(headings)))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, headings):
        rtl_paragraph(cell.paragraphs[0], heading, size_pt=10)
    rtl_paragraph(table.rows[1].cells[0].paragraphs[0], "{%tr for i in institutes %}")
    for cell, field in zip(table.rows[2].cells, fields):
        rtl_paragraph(cell.paragraphs[0], f"{{{{ i.{field} }}}}", size_pt=10)
    rtl_paragraph(table.rows[3].cells[0].paragraphs[0], "{%tr endfor %}")


def build_case_summary(path: Path) -> None:
    """The Step-5 cover sheet: the compiled case's own documents are appended after it (§10.3)."""
    doc = Docx()
    doc.styles["Normal"].font.size = Pt(11)
    _letterhead(doc)
    rtl_paragraph(doc.add_paragraph(), "پوختەی دۆسیەی دابەشکردنی زەوی", size_pt=14)
    _summary_fields_table(doc)
    doc.add_paragraph()
    rtl_paragraph(doc.add_paragraph(), "هەنگاوەکان", size_pt=12)
    _steps_table(doc)
    doc.add_paragraph()
    rtl_paragraph(doc.add_paragraph(), "دەزگاکان", size_pt=12)
    _institutes_table(doc)
    doc.save(path)


BUILDERS = {
    DocumentTemplate.TemplateType.ELIGIBILITY_SINGLE: (
        build_eligibility_single,
        "Placeholder — eligibility letter",
    ),
    DocumentTemplate.TemplateType.PROCESS_LIST: (
        build_process_list,
        "Placeholder — beneficiary list letter",
    ),
    DocumentTemplate.TemplateType.CASE_SUMMARY: (
        build_case_summary,
        "Placeholder — compiled case summary",
    ),
}


class Command(BaseCommand):
    help = "Create placeholder .docx templates for any type that has no active template yet."

    def add_arguments(self, parser):
        parser.add_argument(
            "--replace",
            action="store_true",
            help="Rebuild even for types that already have an active template (retires it).",
        )

    def handle(self, *args, **options):
        actor = User.objects.filter(role=User.Role.ADMIN, is_active=True).first()
        with tempfile.TemporaryDirectory(prefix="tpl-") as work:
            for template_type, (build, name) in BUILDERS.items():
                # Registering a placeholder makes it active, which RETIRES whatever was active
                # before — including the office's real letter. Skip by default so re-running this
                # for one new type cannot silently replace templates the office uploaded.
                if not options["replace"] and DocumentTemplate.objects.filter(
                    template_type=template_type, is_active=True
                ).exists():
                    self.stdout.write(f"skip {template_type}: an active template already exists")
                    continue
                path = Path(work) / f"{template_type}.docx"
                build(path)
                upload = SimpleUploadedFile(path.name, path.read_bytes())
                template = create_template(
                    template_type=template_type, name=name, upload=upload, actor=actor
                )
                self.stdout.write(
                    self.style.SUCCESS(f"active {template_type} template → {template.file_path}")
                )
