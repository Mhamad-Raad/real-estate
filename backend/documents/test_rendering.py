"""Guards the generation chain: docxtpl fill → row loop → headless LibreOffice → PDF (§6.6).

This is the Iteration 3 de-risk. RTL Sorani output is the flagged risk (§13), and every part of
the chain fails quietly: a bad loop tag renders zero rows, a missing font substitutes glyphs.
Needs LibreOffice, so it runs in the container and skips on a bare native checkout.
"""

import shutil
import tempfile
import unittest
from pathlib import Path

from django.conf import settings
from django.test import SimpleTestCase
from docx import Document as Docx
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from pypdf import PdfReader

from .docx_rtl import rtl_paragraph, rtl_run, rtl_table
from .factories import HAS_LIBREOFFICE, NO_LIBREOFFICE_REASON
from .rendering import RenderError, docx_to_pdf


NAME = "هاوبیر أمید کاک عبدالله"
# Assert on a ligature-free fragment: LibreOffice draws "لله" as a single glyph that pypdf cannot
# map back to code points, so an exact compare fails on text that renders perfectly on the page.
NAME_FRAGMENT = "هاوبیر أمید"
LAST_ROW_NAME = "ڕێبوار ڕەشید"
KURDISH_LETTERS = "ڕێزدار ڵ ۆ ێ ژ گ چ پ ڤ"


def build_template(path: Path) -> None:
    """A miniature of the real letter: one placeholder plus a repeating table row."""
    doc = Docx()
    rtl_paragraph(doc.add_paragraph(), f"ناوی سوودمەند: {{{{ full_name }}}}", size_pt=12)
    rtl_paragraph(doc.add_paragraph(), KURDISH_LETTERS, size_pt=12)

    # A {%tr %} tag deletes the row holding it, so the loop markers get rows of their own:
    # header / for-row (deleted) / content row (repeated) / endfor-row (deleted).
    table = rtl_table(doc.add_table(rows=4, cols=3))
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells, ["ز", "ناو", "ناوی دایک"]):
        rtl_paragraph(cell.paragraphs[0], head)
    rtl_paragraph(table.rows[1].cells[0].paragraphs[0], "{%tr for r in rows %}")
    for cell, field in zip(table.rows[2].cells, ["n", "name", "mother"]):
        rtl_paragraph(cell.paragraphs[0], f"{{{{ r.{field} }}}}")
    rtl_paragraph(table.rows[3].cells[0].paragraphs[0], "{%tr endfor %}")
    doc.save(path)


@unittest.skipUnless(HAS_LIBREOFFICE, NO_LIBREOFFICE_REASON)
class DocxToPdfTests(SimpleTestCase):
    def setUp(self):
        self.work = Path(tempfile.mkdtemp(prefix="render-test-"))
        self.addCleanup(shutil.rmtree, self.work, ignore_errors=True)

    def render(self, context: dict) -> str:
        template = self.work / "template.docx"
        build_template(template)
        filled = self.work / "filled.docx"
        tpl = DocxTemplate(template)
        tpl.render(context)
        tpl.save(filled)

        pdf = docx_to_pdf(filled, self.work / "out")
        self.assertTrue(pdf.is_file())
        reader = PdfReader(str(pdf))
        self.assertEqual(len(reader.pages), 1)
        return "".join(page.extract_text() or "" for page in reader.pages)

    def test_fills_placeholders_and_grows_one_row_per_item(self):
        rows = [
            {"n": "١", "name": "نجاة سلام علي", "mother": "سمیرە حەمە"},
            {"n": "٢", "name": NAME, "mother": "گوڵاڵە ئەحمەد"},
            {"n": "٣", "name": LAST_ROW_NAME, "mother": "شێرین مستەفا"},
        ]
        text = self.render({"full_name": NAME, "rows": rows})

        self.assertIn(NAME_FRAGMENT, text)
        # The last row proves the loop expanded rather than rendering the template row once.
        self.assertIn(LAST_ROW_NAME, text)
        # Kurdish-specific letters survive the font pipeline instead of dropping to boxes.
        for letter in "ڕڵۆێڤ":
            self.assertIn(letter, text)
        # Nothing unrendered leaked into a document that goes to a ministry.
        for marker in ("{{", "{%"):
            self.assertNotIn(marker, text)

    def test_empty_row_list_keeps_the_letter_but_drops_every_data_row(self):
        text = self.render({"full_name": NAME, "rows": []})

        self.assertIn(NAME_FRAGMENT, text)
        self.assertNotIn(LAST_ROW_NAME, text)

    def test_missing_input_file_raises_render_error(self):
        with self.assertRaises(RenderError):
            docx_to_pdf(self.work / "nope.docx", self.work / "out")

    def test_failed_render_never_returns_an_earlier_pdf(self):
        """A superseded letter must not come back as if it were freshly generated.

        LibreOffice exits 0 having produced nothing, so without clearing the target first a
        failed regeneration would hand back the previous client data as a success.
        """
        out = self.work / "out"
        out.mkdir()
        stale = out / "letter.pdf"
        stale.write_bytes(b"%PDF-1.4 SUPERSEDED LETTER")
        # Valid path, unreadable content — what a corrupt uploaded template looks like.
        corrupt = self.work / "letter.docx"
        corrupt.write_bytes(b"PK\x03\x04 truncated archive")

        with self.assertRaises(RenderError):
            docx_to_pdf(corrupt, out)
        self.assertFalse(stale.exists(), "the superseded PDF must not survive a failed render")


class RtlHelperTests(SimpleTestCase):
    """The direction flags are write-once: a duplicated w:bidi/w:rtl is invalid OOXML."""

    def test_helpers_are_idempotent(self):
        doc = Docx()
        paragraph = rtl_paragraph(doc.add_paragraph(), "تاقیکردنەوە")
        rtl_paragraph(paragraph)
        self.assertEqual(len(paragraph._p.get_or_add_pPr().findall(qn("w:bidi"))), 1)

        run = paragraph.runs[0]
        rtl_run(run)
        self.assertEqual(len(run._r.get_or_add_rPr().findall(qn("w:rtl"))), 1)

        table = rtl_table(doc.add_table(rows=1, cols=1))
        rtl_table(table)
        self.assertEqual(len(table._tbl.tblPr.findall(qn("w:bidiVisual"))), 1)

    def test_rtl_run_sits_in_schema_order_after_size(self):
        """Out-of-order OOXML is ignored without error — w:rtl must follow sz/szCs."""
        run = rtl_run(Docx().add_paragraph().add_run("x"), size_pt=12)
        tags = [child.tag.split("}")[1] for child in run._r.get_or_add_rPr()]
        self.assertLess(tags.index("sz"), tags.index("rtl"))
